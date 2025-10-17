import pythoncom
import tempfile
import io
from django.http import FileResponse
from rest_framework.views import APIView
from rest_framework.parsers import MultiPartParser, FormParser
from rest_framework.response import Response
from rest_framework import status
from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.shared import Pt
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_LEFT, TA_RIGHT, TA_CENTER, TA_JUSTIFY
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image as ReportLabImage, Table, TableStyle
from reportlab.lib import colors
from xml.sax.saxutils import escape as xml_escape
import re
from PIL import Image as PILImage
from io import BytesIO

def _to_pt(length_or_none, default=0):
    try:
        return float(getattr(length_or_none, "pt", default) or default)
    except Exception:
        return default

def _align_to_ta(al):
    if al is None:
        return TA_LEFT
    return {
        0: TA_LEFT,   # WD_ALIGN_PARAGRAPH.LEFT
        1: TA_CENTER, # CENTER
        2: TA_RIGHT,  # RIGHT
        3: TA_JUSTIFY # JUSTIFY
    }.get(int(al), TA_LEFT)

def _normalize_inline_text(text: str) -> str:
    """
    - Escape XML special chars for ReportLab mini-markup
    - Preserve multiple spaces (convert to &nbsp;)
    - Convert tabs to a few non-breaking spaces
    - Convert newlines to <br/>
    """
    text = xml_escape(text, entities={"'": "&apos;", '"': "&quot;"})
    text = text.replace("\t", "&nbsp;&nbsp;&nbsp;&nbsp;")
    text = re.sub(r"^ +", lambda m: "&nbsp;" * len(m.group(0)), text)
    text = re.sub(r" {2,}", lambda m: "&nbsp;" * len(m.group(0)), text)
    text = text.replace("\n", "<br/>")
    return text

def _extract_images(doc):
    images = []
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            img = rel.target_part
            image_data = img.blob
            img_type = img.content_type.split('/')[1]
            image_stream = io.BytesIO(image_data)
            images.append((img_type, image_stream))
    return images

def _scale_image(img_stream, max_width=439.27559055118115, max_height=685.8897637795277):
    pil_image = PILImage.open(img_stream)
    img_width, img_height = pil_image.size
    scale_width = max_width / img_width
    scale_height = max_height / img_height
    scale_factor = min(scale_width, scale_height)
    new_width = img_width * scale_factor
    new_height = img_height * scale_factor
    img_stream.seek(0)
    return ReportLabImage(img_stream, width=new_width, height=new_height)

class DocxToPdfView(APIView):
    parser_classes = [MultiPartParser, FormParser]

    def post(self, request, *args, **kwargs):
        try:
            file_obj = request.FILES.get('file')
            if not file_obj:
                return Response({"error": "No file Uploaded"}, status=status.HTTP_400_BAD_REQUEST)

            with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as temp_docx:
                for chunk in file_obj.chunks():
                    temp_docx.write(chunk)
                temp_docx.flush()

            doc = Document(temp_docx.name)
            styles = getSampleStyleSheet()

            images = _extract_images(doc)

            with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as temp_pdf:
                pdf_doc = SimpleDocTemplate(temp_pdf.name, pagesize=A4)
                story = []

                # --- PARAGRAPHS ---
                for idx, para in enumerate(doc.paragraphs):
                    pf = para.paragraph_format

                    font_size = None
                    font_family = "Helvetica"
                    font_weight = "normal"
                    
                    def _map_font_name(word_font: str) -> str:
                        if not word_font:
                            return "Helvetica"  # fallback
                        word_font = word_font.lower()
                        if "times" in word_font:
                            return "Times-Roman"
                        if "arial" in word_font:
                            return "Helvetica"
                        if "courier" in word_font:
                            return "Courier"
                        return "Helvetica"  # default


                    for r in para.runs:
                        if r.font and r.font.size:
                            font_size = r.font.size.pt
                        if r.font and r.font.name:
                            font_family =  _map_font_name(r.font.name)
                        if r.font and r.font.bold:
                            font_weight = "bold"

                    if font_size is None:
                        font_size = 12  # Default font size

                    # Adjust line spacing (leading)
                    leading = None
                    ls_rule = getattr(pf, "line_spacing_rule", None)
                    ls_value = getattr(pf, "line_spacing", None)

                    # Handle different line spacing cases:
                    if ls_rule == WD_LINE_SPACING.MULTIPLE and ls_value:
                        leading = font_size * float(ls_value)  # Multiply font size by specified multiple
                    elif ls_rule == WD_LINE_SPACING.EXACTLY and ls_value:
                        leading = float(ls_value)  # Use exact points for line spacing
                    elif ls_rule == WD_LINE_SPACING.AT_LEAST and ls_value:
                        leading = max(font_size, float(ls_value))  # At least a given value
                    elif ls_rule == WD_LINE_SPACING.SINGLE:
                        leading = font_size  # Single spacing
                    elif ls_rule == WD_LINE_SPACING.ONE_POINT_FIVE:
                        leading = font_size * 1.5  # 1.5x spacing
                    elif ls_rule == WD_LINE_SPACING.DOUBLE:
                        leading = font_size * 2  # Double spacing
                    else:
                        leading = font_size * 1.2  # Default to 1.2x font size for undefined rules

                    # Paragraph spacing / indents / alignment
                    space_before = _to_pt(getattr(pf, "space_before", None), 0)
                    space_after = _to_pt(getattr(pf, "space_after", None), 0)
                    left_indent = _to_pt(getattr(pf, "left_indent", None), 0)
                    right_indent = _to_pt(getattr(pf, "right_indent", None), 0)
                    first_line = _to_pt(getattr(pf, "first_line_indent", None), 0)
                    alignment = _align_to_ta(getattr(para, "alignment", None))

                    pstyle = ParagraphStyle(
                        name=f"p{idx}",
                        parent=styles["Normal"],
                        fontName=font_family,
                        fontSize=font_size,
                        leading=leading,
                        spaceBefore=space_before,
                        spaceAfter=space_after,
                        leftIndent=left_indent,
                        rightIndent=right_indent,
                        firstLineIndent=first_line,
                        alignment=alignment,
                        fontWeight=font_weight
                    )

                    formatted_text = ""
                    if para.runs:
                        for run in para.runs:
                            if not run.text:
                                continue
                            txt = _normalize_inline_text(run.text)

                            if run.bold:
                                txt = f"<b>{txt}</b>"
                            if run.italic:
                                txt = f"<i>{txt}</i>"
                            if run.underline:
                                txt = f"<u>{txt}</u>"

                            formatted_text += txt
                    else:
                        formatted_text = "&nbsp;"

                    sname = (getattr(para.style, "name", "") or "").lower()
                    if "bullet" in sname or "list" in sname:
                        if first_line == 0 and left_indent == 0:
                            pstyle.firstLineIndent = -10
                            pstyle.leftIndent = 20
                        formatted_text = f"•&nbsp;{formatted_text}"

                    story.append(Paragraph(formatted_text if formatted_text.strip() else "&nbsp;", pstyle))

                # --- ADD IMAGES ---
                for img_type, img_stream in images:
                    reportlab_img = _scale_image(img_stream)
                    story.append(reportlab_img)

                # Account for trailing whitespace between the last paragraph and the next block
                story.append(Spacer(1, 0))

                # --- TABLES (simple rendering; optional but nicer spacing than rows of Paragraphs) ---
                for table in doc.tables:
                    data = []
                    for row in table.rows:
                        data.append([cell.text if cell.text else "" for cell in row.cells])

                    t = Table(data, hAlign="LEFT")
                    t.setStyle(TableStyle([
                        ("FONT", (0, 0), (-1, -1), "Helvetica", 10),
                        ("GRID", (0, 0), (-1, -1), 0.25, colors.black),
                        ("VALIGN", (0, 0), (-1, -1), "TOP"),
                        ("LEFTPADDING", (0, 0), (-1, -1), 4),
                        ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                        ("TOPPADDING", (0, 0), (-1, -1), 2),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 2),
                    ]))
                    story.append(Spacer(1, 6))
                    story.append(t)
                    story.append(Spacer(1, 6))

                pdf_doc.build(story)

                return FileResponse(
                    open(temp_pdf.name, "rb"),
                    as_attachment=True,
                    filename="pdfFile.pdf",
                )
        except Exception as e:
            return Response({"error": str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
